"""
DOCX report generator module with Norstella styling and one-page limit.
"""
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple, Union
import datetime
import re
import io
import logging

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips
from docx.shared import Length

from src.config import TEMPLATE_PATH

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates formatted one-page DOCX reports with Norstella styling."""
    
    NORSTELLA_BLUE = RGBColor(31, 73, 125)  # Primary brand color
    NORSTELLA_GRAY = RGBColor(68, 84, 106)  # Secondary color
    
    # Page size constants (in inches)
    PAGE_WIDTH = 8.5
    PAGE_HEIGHT = 11.0
    MARGIN = 1.0
    CONTENT_WIDTH = PAGE_WIDTH - (2 * MARGIN)
    CONTENT_HEIGHT = PAGE_HEIGHT - (2 * MARGIN)
    
    def __init__(self, template_path: Optional[str] = None):
        """Initialize report generator."""
        self.doc = Document()
        self._setup_norstella_styles()
        self._setup_page_layout()
        self.content_height = 0  # Track content height

    def _setup_page_layout(self):
        """Set up page layout for one-page format."""
        section = self.doc.sections[0]
        section.page_height = Inches(self.PAGE_HEIGHT)
        section.page_width = Inches(self.PAGE_WIDTH)
        section.left_margin = Inches(self.MARGIN)
        section.right_margin = Inches(self.MARGIN)
        section.top_margin = Inches(self.MARGIN)
        section.bottom_margin = Inches(self.MARGIN)

    def _setup_norstella_styles(self):
        """Set up document styles for one-page format."""
        # Title style (smaller than before)
        title_style = self.doc.styles.add_style('NorstTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Segoe UI'  # type: ignore
        title_style.font.size = Pt(18)  # Reduced from 24  # type: ignore
        title_style.font.bold = True  # type: ignore
        title_style.font.color.rgb = self.NORSTELLA_BLUE  # type: ignore
        title_style.paragraph_format.space_before = Pt(12)  # type: ignore
        title_style.paragraph_format.space_after = Pt(6)  # type: ignore
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore
        
        # Subtitle style (smaller)
        subtitle_style = self.doc.styles.add_style('NorstSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = 'Segoe UI'  # type: ignore
        subtitle_style.font.size = Pt(12)  # Reduced from 16  # type: ignore
        subtitle_style.font.color.rgb = self.NORSTELLA_GRAY  # type: ignore
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore
        subtitle_style.paragraph_format.space_after = Pt(12)  # type: ignore
        
        # Section header style (smaller)
        header_style = self.doc.styles.add_style('NorstHeader', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = 'Segoe UI'  # type: ignore
        header_style.font.size = Pt(11)  # Reduced from 14  # type: ignore
        header_style.font.bold = True  # type: ignore
        header_style.font.color.rgb = self.NORSTELLA_BLUE  # type: ignore
        header_style.paragraph_format.space_before = Pt(6)  # type: ignore
        header_style.paragraph_format.space_after = Pt(3)  # type: ignore
        
        # Body style (compact)
        body_style = self.doc.styles.add_style('NorstBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'  # type: ignore
        body_style.font.size = Pt(10)  # Reduced from 11  # type: ignore
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # type: ignore
        body_style.paragraph_format.line_spacing = 1.0  # Single spacing  # type: ignore
        body_style.paragraph_format.space_after = Pt(3)  # type: ignore
        
        # List style (compact)
        list_style = self.doc.styles.add_style('NorstList', WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.name = 'Calibri'  # type: ignore
        list_style.font.size = Pt(10)  # type: ignore
        list_style.paragraph_format.line_spacing = 1.0  # type: ignore
        list_style.paragraph_format.left_indent = Inches(0.15)  # type: ignore
        list_style.paragraph_format.space_after = Pt(2)  # type: ignore

    def _parse_markdown_text(self, text: str, paragraph):
        """
        Parse markdown-formatted text and add it to a paragraph with proper Word formatting.
        
        Args:
            text: Text containing markdown formatting
            paragraph: docx paragraph object to add formatted text to
        """
        # Split text by markdown patterns while preserving the delimiters
        # This regex captures bold (**text**), italic (*text*), and regular text
        pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part.strip():
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                bold_text = part[2:-2]  # Remove ** markers
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text  
                italic_text = part[1:-1]  # Remove * markers
                run = paragraph.add_run(italic_text)
                run.italic = True
            else:
                # Regular text
                paragraph.add_run(part)

    def _add_formatted_paragraph(self, text: str, style: str = 'NorstBody'):
        """
        Add a paragraph with markdown formatting converted to Word formatting.
        
        Args:
            text: Text that may contain markdown formatting
            style: Paragraph style to apply
        """
        # Clean up the text first
        text = text.strip()
        if not text:
            return
        
        # Create the paragraph
        paragraph = self.doc.add_paragraph(style=style)
        
        # Parse and add the formatted text
        self._parse_markdown_text(text, paragraph)

    def _process_section_content(self, content: str) -> List[str]:
        """
        Process section content and split into paragraphs, handling markdown headers.
        
        Args:
            content: Raw content text with possible markdown formatting
            
        Returns:
            List of processed paragraphs with (type, text) tuples
        """
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        processed_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # Check if this is a markdown header (starts with **)
            if para.startswith('**') and '**' in para[2:]:
                # This is a subheader - extract the header text
                header_match = re.match(r'\*\*(.*?)\*\*(.*)', para, re.DOTALL)
                if header_match:
                    header_text = header_match.group(1).strip()
                    remaining_text = header_match.group(2).strip()
                    
                    # Only add header if there's meaningful content following it
                    if header_text and remaining_text and len(remaining_text) > 20:
                        processed_paragraphs.append(('header', header_text))
                        processed_paragraphs.append(('body', remaining_text))
                    elif header_text and remaining_text:
                        # If there's a header but minimal content, combine them
                        combined_text = f"{header_text}: {remaining_text}"
                        processed_paragraphs.append(('body', combined_text))
                    elif remaining_text:
                        # If there's only content, treat as regular paragraph
                        processed_paragraphs.append(('body', remaining_text))
                else:
                    # Fallback - treat as regular text
                    processed_paragraphs.append(('body', para))
            else:
                # Regular paragraph - only add if it has substantial content
                if len(para) > 10:  # Minimum length threshold
                    processed_paragraphs.append(('body', para))
        
        return processed_paragraphs

    def add_section(self, title: str, content: str, max_length: Optional[int] = None):
        """
        Add a section with proper Word formatting instead of markdown.
        
        Args:
            title: Section title
            content: Section content text (may contain markdown)
            max_length: Maximum number of characters (optional)
        """
        # Skip sections with no meaningful content
        if not content or len(content.strip()) < 20:
            return
        
        # Add section header
        header = self.doc.add_paragraph(title, style='NorstHeader')
        
        # Truncate content if needed
        if max_length and len(content) > max_length:
            content = content[:max_length-3] + "..."
        
        # Process the content to handle markdown formatting
        processed_content = self._process_section_content(content)
        
        # Only add section if there's actual content to display
        if not processed_content:
            return
        
        for content_type, text in processed_content:
            if content_type == 'header':
                # Add as a sub-header with bold formatting
                sub_header = self.doc.add_paragraph(style='NorstBody')
                run = sub_header.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
            elif content_type == 'body':
                # Check if it's a bullet point
                if text.strip().startswith(('•', '-', '*')):
                    # Handle bullet points
                    bullet_text = text.strip()[1:].strip()
                    if bullet_text:  # Only add if there's actual content
                        p = self.doc.add_paragraph(style='NorstList')
                        p.add_run(bullet_text)
                else:
                    # Regular paragraph
                    self._add_formatted_paragraph(text)

    def add_cover_page(self, title: str, subtitle: Optional[str] = None):
        """Add a cover page to the document."""
        # Add title
        title_para = self.doc.add_paragraph(title, style='NorstTitle')
        
        # Add subtitle if provided
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle, style='NorstSubtitle')
    
    def add_norstella_header(self, company_name: str, report_type: str = "Flash Report"):
        """Add Norstella header with logo and report title."""
        # Create a table for the header layout
        header_table = self.doc.add_table(rows=1, cols=3)
        header_table.style = 'Table Grid'
        header_table.autofit = False
        
        # Set column widths (logo, title, empty)
        header_table.columns[0].width = Inches(1.5)  # Logo column
        header_table.columns[1].width = Inches(5.5)  # Title column
        header_table.columns[2].width = Inches(1.5)  # Empty column
        
        # Add logo to first cell
        logo_cell = header_table.cell(0, 0)
        try:
            logo_path = TEMPLATE_PATH / "Images" / "Norstella_color_positive_RGB_(2).png"
            if logo_path.exists():
                logo_para = logo_cell.paragraphs[0]
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Inches(1.2))
        except Exception as e:
            logger.warning(f"Could not add logo: {e}")
            logo_cell.text = "NORSTELLA"
        
        # Add title to center cell
        title_cell = header_table.cell(0, 1)
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company name and report type
        company_run = title_para.add_run(f"{company_name} {report_type}")
        company_run.font.name = 'Segoe UI'
        company_run.font.size = Pt(16)
        company_run.font.bold = True
        company_run.font.color.rgb = self.NORSTELLA_BLUE
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def add_three_column_layout(self, executive_summary: str, key_takeaways: str, financial_highlights: str):
        """Add three-column layout matching the IQVIA report format."""
        # Create main table for three-column layout
        main_table = self.doc.add_table(rows=2, cols=2)
        main_table.style = 'Table Grid'
        main_table.autofit = False
        
        # Set column widths
        main_table.columns[0].width = Inches(4.0)  # Left column (Executive Summary)
        main_table.columns[1].width = Inches(4.0)  # Right column (Key Takeaways)
        
        # Set row heights
        main_table.rows[0].height = Inches(4.0)  # Top row
        main_table.rows[1].height = Inches(3.0)  # Bottom row (Financial Highlights)
        
        # Executive Summary (top left)
        exec_cell = main_table.cell(0, 0)
        exec_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.TOP  # type: ignore
        exec_para = exec_cell.paragraphs[0]
        exec_run = exec_para.add_run("EXECUTIVE SUMMARY")
        exec_run.font.name = 'Segoe UI'  # type: ignore
        exec_run.font.size = Pt(12)  # type: ignore
        exec_run.font.bold = True  # type: ignore
        exec_run.font.color.rgb = self.NORSTELLA_BLUE  # type: ignore
        
        # Add executive summary content
        exec_content_para = exec_cell.add_paragraph()
        exec_content_para.style = 'NorstBody'
        self._add_formatted_text_to_paragraph(executive_summary, exec_content_para)
        
        # Key Takeaways (top right)
        key_cell = main_table.cell(0, 1)
        key_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.TOP  # type: ignore
        key_para = key_cell.paragraphs[0]
        key_run = key_para.add_run("KEY TAKEAWAYS")
        key_run.font.name = 'Segoe UI'  # type: ignore
        key_run.font.size = Pt(12)  # type: ignore
        key_run.font.bold = True  # type: ignore
        key_run.font.color.rgb = self.NORSTELLA_BLUE  # type: ignore
        
        # Add key takeaways content
        key_content_para = key_cell.add_paragraph()
        key_content_para.style = 'NorstBody'
        self._add_formatted_text_to_paragraph(key_takeaways, key_content_para)
        
        # Financial Highlights (bottom, spanning both columns)
        financial_cell = main_table.cell(1, 0)
        financial_cell.merge(main_table.cell(1, 1))  # Merge bottom row
        financial_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.TOP  # type: ignore
        
        financial_para = financial_cell.paragraphs[0]
        financial_run = financial_para.add_run("FINANCIAL HIGHLIGHTS")
        financial_run.font.name = 'Segoe UI'
        financial_run.font.size = Pt(12)
        financial_run.font.bold = True
        financial_run.font.color.rgb = self.NORSTELLA_BLUE
        
        # Add financial highlights content
        financial_content_para = financial_cell.add_paragraph()
        financial_content_para.style = 'NorstBody'
        self._add_formatted_text_to_paragraph(financial_highlights, financial_content_para)
    
    def _add_formatted_text_to_paragraph(self, text: str, paragraph):
        """Add formatted text to an existing paragraph."""
        # Clean up the text
        text = text.strip()
        if not text:
            return
        
        # Parse and add the formatted text
        self._parse_markdown_text(text, paragraph)
    
    def add_footer(self):
        """Add footer with confidential notice and footnotes."""
        # Add spacing
        self.doc.add_paragraph()
        
        # Add confidential notice
        confidential_para = self.doc.add_paragraph()
        confidential_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidential_run = confidential_para.add_run("CONFIDENTIAL AND PROPRIETARY")
        confidential_run.font.name = 'Segoe UI'
        confidential_run.font.size = Pt(8)
        confidential_run.font.color.rgb = self.NORSTELLA_GRAY
        
        # Add footnotes section
        footnotes_para = self.doc.add_paragraph()
        footnotes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footnotes_run = footnotes_para.add_run("Footnotes:")
        footnotes_run.font.name = 'Segoe UI'
        footnotes_run.font.size = Pt(8)
        footnotes_run.font.bold = True
        footnotes_run.font.color.rgb = self.NORSTELLA_GRAY
        
        # Add sample footnotes (these would be dynamic based on content)
        footnote1 = self.doc.add_paragraph()
        footnote1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footnote1_run = footnote1.add_run("* Technology & Analytics Solutions")
        footnote1_run.font.name = 'Segoe UI'
        footnote1_run.font.size = Pt(8)
        footnote1_run.font.color.rgb = self.NORSTELLA_GRAY
        
        footnote2 = self.doc.add_paragraph()
        footnote2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footnote2_run = footnote2.add_run("** Research & Development Solutions")
        footnote2_run.font.name = 'Segoe UI'
        footnote2_run.font.size = Pt(8)
        footnote2_run.font.color.rgb = self.NORSTELLA_GRAY
        
        footnote3 = self.doc.add_paragraph()
        footnote3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footnote3_run = footnote3.add_run("*** Contract Sales & Medical Solutions")
        footnote3_run.font.name = 'Segoe UI'
        footnote3_run.font.size = Pt(8)
        footnote3_run.font.color.rgb = self.NORSTELLA_GRAY

    def save(self, filename: str):
        """Save the document to a file."""
        self.doc.save(filename)
    
    def get_docx_bytes(self) -> bytes:
        """
        Get the document as bytes for Streamlit download_button.
        This is the recommended approach for Streamlit Cloud compatibility.
        
        Returns:
            bytes: The complete docx document as bytes
        """
        try:
            bio = io.BytesIO()
            self.doc.save(bio)
            bio.seek(0)
            document_bytes = bio.getvalue()
            
            # Validate that we have actual content
            if len(document_bytes) < 1000:  # A valid docx should be at least 1KB
                logger.warning("Generated docx appears to be too small, may be corrupted")
            
            logger.info(f"Generated docx document: {len(document_bytes)} bytes")
            return document_bytes
            
        except Exception as e:
            logger.error(f"Error generating docx bytes: {str(e)}")
            # Create a minimal fallback document
            fallback_doc = Document()
            fallback_doc.add_paragraph("Error generating report. Please try again.")
            fallback_bio = io.BytesIO()
            fallback_doc.save(fallback_bio)
            fallback_bio.seek(0)
            return fallback_bio.getvalue()

    def add_formatted_text(self, text: str, style: str = 'NorstBody'):
        """Add formatted text to the document."""
        self._add_formatted_paragraph(text, style) 